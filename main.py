import os
import sys
import time

import yara
from PyQt6.QtCore import pyqtSignal, Qt
from PyQt6.QtWidgets import QWidget, QApplication, QMainWindow, QPushButton, QLabel, QMessageBox, QVBoxLayout, \
    QFileDialog, QHBoxLayout, QProgressBar, QLineEdit, QTextEdit
from docx import Document


def create_rules_list(address):
    yara_list = {}
    files = os.listdir(address)
    for file_name in files:
        if '.yar' in file_name:
            yara_list[file_name.split('.yar')[0]] = address + file_name
    return yara_list


def print_in_document(path_to_check, matches):
    # print(f"The file {path_to_check} follows the following rules:")
    count = 0
    for match in matches:

        print("\t", match)
        count += 1

        if os.path.exists(f'docxDir/{str(match)}.docx'):
            doc_to_write = Document(f'docxDir/{str(match)}.docx')
            doc_to_write.add_paragraph(f"The file {path_to_check}")
            doc_to_write.save(f'docxDir/{str(match)}.docx')
        else:
            document = Document()
            document.add_paragraph(f"{str(match)}:")
            document.add_paragraph(f"The file {path_to_check}")
            document.save(f'docxDir/{str(match)}.docx')

    return count


def checking_files_in_directory(rules_for_checking, address):
    directory_path = address
    x = 0
    for filename in os.listdir(directory_path):
        filepath = os.path.join(directory_path, filename)
        if os.path.isfile(filepath):
            matches_files = rules_for_checking.match(filepath)
            if matches_files:
                x = print_in_document(filename, matches_files)

    return x


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        _signal = pyqtSignal(int)

        self.dir_list = None
        self.filename = None
        self.plainTextEdit = None
        self.text_to_save = None
        self.setWindowTitle("AntiVirus Demo")
        self.setFixedSize(410, 300)

        self.message_of_checking_exist_file = QMessageBox()

        layout_first_card = QHBoxLayout()
        layout_second_card = QVBoxLayout()
        layout_third_card = QVBoxLayout()
        layout_all = QVBoxLayout()

        # description of layout_first_card ----------------------------------------
        result_of_scanning = QLabel("Выберите")

        self.btn_to_choose_file = QPushButton("Файл")
        self.btn_to_choose_file.setCheckable(True)
        self.btn_to_choose_file.clicked.connect(self.the_btn_to_choose_file_was_clicked)
        or_ = QLabel("или")
        or_.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.btn_to_choose_dir = QPushButton("Папка")
        self.btn_to_choose_dir.setCheckable(True)
        self.btn_to_choose_dir.clicked.connect(self.the_btn_to_choose_dir_was_clicked)
        end_of_sent = QLabel("для сканирования")

        layout_first_card.addWidget(result_of_scanning)
        layout_first_card.addWidget(self.btn_to_choose_file)
        layout_first_card.addWidget(or_)
        layout_first_card.addWidget(self.btn_to_choose_dir)
        layout_first_card.addWidget(end_of_sent)
        # -------------------------------------------------------------------------

        # description of layout_second_card ----------------------------------------
        self.non_editable_line_edit = QLineEdit(self)
        self.non_editable_line_edit.setReadOnly(True)
        # self.non_editable_line_edit.resize(150, 250)

        self.pbar = QProgressBar(self)
        self.pbar.setValue(0)
        self.pbar.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.pbar.setStyleSheet("QProgressBar::chunk "
                  "{"
                    "background-color: pink;"
                  "}")

        self.btn_to_scan = QPushButton("Сканировать")
        self.btn_to_scan.setCheckable(True)
        self.btn_to_scan.clicked.connect(self.the_button_was_clicked)

        layout_second_card.addWidget(self.non_editable_line_edit)
        layout_second_card.addWidget(self.pbar)
        layout_second_card.addWidget(self.btn_to_scan)
        # --------------------------------------------------------------------------

        # description of layout_third_card -----------------------------------------
        result_all = QLabel("Результаты сканирования:")
        result_all.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.result = QTextEdit(self)
        self.result.resize(150, 100)
        self.result.setReadOnly(True)

        layout_third_card.addWidget(result_all)
        layout_third_card.addWidget(self.result)

        # --------------------------------------------------------------------------

        layout_all.addLayout(layout_first_card)
        layout_all.addLayout(layout_second_card)
        layout_all.addLayout(layout_third_card)

        container = QWidget()

        container.setLayout(layout_all)

        self.setCentralWidget(container)

    def the_btn_to_choose_file_was_clicked(self):
        filename, filetype = QFileDialog.getOpenFileName(self,
                                                         "Выбрать файл",
                                                         ".",
                                                         "All Files(*);;Text Files(*.txt);;JPEG Files(*.jpeg);;\
                                                         PNG Files(*.png);;GIF File(*.gif)")
        if filename != '':
            # self.btn_to_choose_dir.setEnabled(False)
            self.non_editable_line_edit.setText(filename)
            self.filename = filename

    def the_btn_to_choose_dir_was_clicked(self):  # <-----
        dir_list = QFileDialog.getExistingDirectory(self, "Выбрать папку", ".")
        if dir_list != '':
            # self.btn_to_choose_file.setEnabled(False)
            self.non_editable_line_edit.setText(dir_list)
            # print(self.non_editable_line_edit.text())
            self.dir_list = dir_list

    def the_button_was_clicked(self):

        rules = yara.compile(filepaths=create_rules_list('rulesDir/'))

        for i in range(101):
            # slowing down the loop
            time.sleep(0.01)

            # setting value to progress bar
            self.pbar.setValue(i)

        file_or_dir_to_scan = self.non_editable_line_edit.text()
        print(file_or_dir_to_scan)
        if os.path.isdir(file_or_dir_to_scan):
            count = checking_files_in_directory(rules, self.dir_list)
            self.result.setText(f"Просканирована директория: {self.dir_list} . Выявлено угроз: {count}.")
        else:
            matches = rules.match(self.filename)
            print(matches.values())
            filename = os.path.basename(self.filename).split('.')[0]
            count = print_in_document(filename, matches)
            self.result.setText(f"Просканирован файл: {filename}. Выявлено угроз: {count}.")


app = QApplication(sys.argv)
window = MainWindow()
window.show()

app.exec()
